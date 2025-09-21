import io
import re
import pdfplumber
import docx
import pytesseract
import fitz
from PIL import Image
from transformers import pipeline
from flask import Flask, request, jsonify, send_file
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import json
import os
from datetime import datetime
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables from .env file in the same directory as this script
script_dir = os.path.dirname(os.path.abspath(__file__))
env_path = os.path.join(script_dir, '.env')
load_dotenv(dotenv_path=env_path)

# Debug: Print loaded environment variables
print(f"Script directory: {script_dir}")
print(f"Env file path: {env_path}")
print(f"Env file exists: {os.path.exists(env_path)}")
gemini_key = os.environ.get('GEMINI_API_KEY', 'Not found')
print(f"GEMINI_API_KEY: {gemini_key}")

# Flask app
app = Flask(__name__)

# Load environment variables
GOOGLE_CLOUD_PROJECT = os.environ.get("GOOGLE_CLOUD_PROJECT", "legalklarity")
GOOGLE_CLOUD_LOCATION = os.environ.get("GOOGLE_CLOUD_LOCATION", "us-central1")
GOOGLE_APPLICATION_CREDENTIALS = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

# Initialize Gemini AI if API key is available
GEMINI_AVAILABLE = False
if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        GEMINI_AVAILABLE = True
        print(f"Gemini API configured successfully with key: {GEMINI_API_KEY[:10]}...")  # Print first 10 chars for debugging
    except Exception as e:
        print(f"Failed to initialize Gemini AI: {e}")
else:
    print("No GEMINI_API_KEY found in environment variables")

_classifier = None
def get_classifier():
    global _classifier
    if _classifier is None:
        _classifier = pipeline(
            "zero-shot-classification",
            model="valhalla/distilbart-mnli-12-3",
            device=-1  # CPU
        )
    return _classifier

# Section cues to check for agreements

POSITIVE_LABELS = [
    "agreement", "legal contract", "rental agreement", "lease agreement",
    "service agreement", "tenant-landlord agreement", "terms and conditions",
    "offer letter", "internship agreement", "employment contract",
    "student agreement", "job offer", "internship terms"
]


SECTION_CUES = [
    "agreement", "security deposit", "rental period", "payment terms",
    "termination", "arbitration", "jurisdiction",
    "witness", "signatory", "governing law", "parties", "definitions",
    "probation period", "internship duration", "performance", 
    "salary", "compensation", "notice period", "work expectations",
    "attendance", "leaves", "certificate", "offer letter"
]

# Helpers
def safe_join_text(parts):
    return "\n".join([p for p in parts if p])

def chunk_text(text, max_words=300, max_chunks=10):
    """
    Split text into chunks of specified word count
    
    Args:
        text (str): Text to chunk
        max_words (int): Maximum words per chunk
        max_chunks (int): Maximum number of chunks to return
        
    Returns:
        list: List of text chunks
    """
    words = text.split()
    chunks = [" ".join(words[i:i + max_words]) for i in range(0, len(words), max_words)]
    return chunks[:max_chunks]

def heuristic_score(text):
    t = (text or "").lower()
    found = sum(1 for k in SECTION_CUES if re.search(r"\b" + re.escape(k) + r"\b", t))
    return found / max(1, len(SECTION_CUES))

def classify_agreement(text):
    details = {
        "chunks": 0,
        "votes": 0,
        "vote_ratio": 0.0,
        "heuristic": 0.0,
        "avg_chunk_score": 0.0,
        "reason": ""
    }
    if not text.strip():
        details["reason"] = "empty_text"
        return False, details
    chunks = chunk_text(text, max_words=300, max_chunks=10)
    details["chunks"] = len(chunks)
    votes, per_chunk_scores = 0, []
    CHUNK_THRESHOLD = 0.5
    for ch in chunks:
        try:
            classifier = get_classifier()
            res = classifier(
                ch,
                candidate_labels=POSITIVE_LABELS,
                multi_label=True,
                hypothesis_template="This text is about {}."
            )
            best = max(res["scores"]) if res["scores"] else 0.0
        except Exception as e:
            print(f"Model error: {e}")
            best = 0.0
        per_chunk_scores.append(best)
        if best >= CHUNK_THRESHOLD:
            votes += 1
    ratio = votes / len(chunks)
    heur = heuristic_score(text)
    details.update({
        "votes": votes,
        "vote_ratio": round(ratio, 3),
        "heuristic": round(heur, 3),
        "avg_chunk_score": round(sum(per_chunk_scores) / max(1, len(per_chunk_scores)), 3)
    })
    accept = (ratio >= 0.4) or (heur >= 0.4)
    if not accept:
        details["reason"] = "low_confidence"
    return accept, details

# Enhanced document analysis function using Legal-Document-Analyzer approach
def analyze_legal_document(text, document_type="legal"):
    """
    Analyze a legal document and provide structured insights using Legal-Document-Analyzer approach.
    Returns a dictionary with summary, key terms, clauses, risks,
    recommendations, and critical legal metadata.
    
    Args:
        text (str): Extracted text from legal document
        document_type (str): Type of document (default: "legal")
    
    Returns:
        dict: Structured analysis with 12 categories
    """
    
    # If Gemini is not available, use fallback analysis
    if not GEMINI_AVAILABLE:
        return create_fallback_analysis(text, document_type)
    
    # Enhanced prompt from Legal-Document-Analyzer approach
    prompt = f"""
    You are an expert legal document analyst reviewing a {document_type}.

    Each field must be filled as defined below:

    1. summary: A 2-3 sentence plain English summary of the document. Avoid legal jargon.  
    2. key_terms: A list of important legal terms or phrases with simple, non-legal definitions (e.g., "Indemnity": "One party promises to cover losses if something goes wrong").  
    3. main_clauses: A list of the main sections/clauses in the document, explained briefly in plain English.  
    4. risks: Any potential risks, unfair obligations, or concerning clauses for the user.  
    5. recommendations: Practical suggestions for what the user should pay attention to or clarify.  
    6. parties: A list naming and briefly describing the main parties involved. If unclear, return [].  
    7. jurisdiction: The governing law or legal jurisdiction. If not found, return "".
    8. obligations: A list of key duties or actions each party must perform (Do Not search for specific obligations, only general duties or actions).  
    9. critical_dates: Important dates, deadlines, or timeframes mentioned.  
    10. missing_or_unusual: Clauses that are missing, unusual, or raise concerns (e.g., "No termination clause found").  
    11. compliance_issues: Any explicit compliance or regulatory issues (e.g., data privacy, employment law, consumer protection).  
    12. next_steps: Action items or questions the user should ask before agreeing/signing.

    Analyze the document and respond ONLY in valid JSON with these fields:

    {{
    "summary": "2-3 sentence plain English summary (avoid legal jargon).",
    "key_terms": ["term1: definition1", "term2: definition2", ...],
    "main_clauses": ["main clause 1", "main clause 2", ...],
    "risks": ["notable risk 1", "notable risk 2", ...],
    "recommendations": ["practical recommendation 1", "practical recommendation 2", ...],
    "parties": ["Party 1: role", "Party 2: role", ...],
    "jurisdiction": "Legal jurisdiction or governing law, or empty if not stated, ...",
    "obligations": ["**obligor1**: obligation 1", "**obligor2**: obligation 2", ..., "**obligee 1**: obligation 1", "**obligee 2**: obligation 2", ...],
    "critical_dates": ["important date or deadline", ...],
    "missing_or_unusual": ["clause that is missing/unusual", ...],
    "compliance_issues": ["compliance or regulatory concern", ...],
    "next_steps": ["question to ask lawyer", "action before signing"]
    }}

    Rules:
    - If information is not found, return an empty string ("") or empty list ([]).
    - Do not include explanations outside the JSON.
    - Keep responses concise and accessible for non-lawyers.
    - While generating each field, add more and more explanation and context to each field.
    - Strictly return JSON with the specified fields and no additional fields.

    Document: {text[:30000]}  # Limit to prevent token overflow
    """

    try:
        # Initialize Gemini model
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Generate response
        response = model.generate_content(prompt)
        
        # Parse and validate JSON response
        analysis = parse_gemini_response(response.text)
        return analysis
        
    except json.JSONDecodeError as e:
        # Fallback to basic analysis if JSON parsing fails
        return create_fallback_analysis(text, document_type)
    except Exception as e:
        # Return error structure
        return {
            "error": f"Analysis failed: {str(e)}",
            "summary": "Document analysis could not be completed due to technical issues.",
            "key_terms": [],
            "main_clauses": [],
            "critical_dates": [],
            "parties": [],
            "jurisdiction": "Not available",
            "obligations": [],
            "risks": [],
            "recommendations": [],
            "missing_or_unusual": [],
            "compliance_issues": [],
            "next_steps": []
        }

# Response parsing function from Legal-Document-Analyzer approach
def parse_gemini_response(response: str) -> dict:
    """Parse the response from Gemini and return a dictionary."""
    try:
        # Check if the response is empty
        if not response.strip():
            return {}

        # Find the first occurrence of the opening curly bracket (`{`)
        start_index = response.find("{")
        if start_index == -1:
            return {}

        # Find the last occurrence of the closing curly bracket (`}`)
        end_index = response.rfind("}")
        if end_index == -1:
            return {}

        # Ensure the start_index is before the end_index
        if start_index > end_index:
            return {}

        # Trim the response to include only the valid JSON
        trimmed_response = response[start_index:end_index + 1]

        # Parse the JSON response
        parsed_data = json.loads(trimmed_response)
        return parsed_data
    except json.JSONDecodeError as e:
        # Fallback in case of invalid JSON
        return {"error": f"Failed to parse response: {str(e)}", "raw": response}

# Document type detection
def detect_document_type(text):
    """
    Enhanced document type detection
    
    Returns:
        str: Detected document type
    """
    text_lower = text.lower()
    
    # Document type patterns
    patterns = {
        "rental agreement": ["rent", "lease", "tenant", "landlord", "security deposit"],
        "employment contract": ["employment", "employee", "employer", "salary", "position"],
        "service agreement": ["service", "provider", "client", "deliverable"],
        "loan agreement": ["loan", "borrower", "lender", "interest rate"],
        "nda": ["confidential", "non-disclosure", "secrecy"],
        "purchase agreement": ["purchase", "buy", "sell", "buyer", "seller"],
        "internship agreement": ["internship", "intern", "supervisor", "internship period"]
    }
    
    # Score each document type
    scores = {}
    for doc_type, keywords in patterns.items():
        score = sum(1 for keyword in keywords if keyword in text_lower)
        scores[doc_type] = score
    
    # Return highest scoring document type
    if scores:
        best_match = max(scores.items(), key=lambda x: x[1])
        if best_match[1] > 0:
            return best_match[0]
    
    return "general legal document"

# Fallback analysis function
def create_fallback_analysis(text, document_type):
    """
    Create basic analysis when AI analysis fails
    
    Returns:
        dict: Basic analysis structure
    """
    # Extract first few sentences for summary
    sentences = text.split('.')
    summary = '. '.join(sentences[:3]) + '.' if len(sentences) > 3 else text[:500]
    
    return {
        "summary": summary,
        "key_terms": [],
        "main_clauses": [],
        "critical_dates": [],
        "parties": [],
        "jurisdiction": "Not analyzed",
        "obligations": [],
        "risks": [],
        "recommendations": ["Have a legal professional review this document"],
        "missing_clauses": [],
        "compliance_issues": [],
        "next_steps": ["Review document with legal counsel"]
    }

# File extraction functions
def extract_pdf(file_stream):
    try:
        file_stream.seek(0)
        with pdfplumber.open(file_stream) as pdf:
            return safe_join_text([p.extract_text() for p in pdf.pages])
    except Exception as e:
        print(f"PDF extract error: {e}")
        try:
            file_stream.seek(0)
            doc = fitz.open(stream=file_stream.read(), filetype="pdf")
            texts = []
            for p in doc:
                pix = p.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                texts.append(pytesseract.image_to_string(img))
            return "\n".join(texts)
        except Exception as e2:
            print(f"PDF OCR error: {e2}")
            return ""

def extract_docx(file_stream):
    try:
        file_stream.seek(0)
        doc = docx.Document(io.BytesIO(file_stream.read()))
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    except Exception as e:
        print(f"DOCX extract error: {e}")
        return ""

def extract_image(file_stream):
    try:
        file_stream.seek(0)
        img = Image.open(file_stream).convert("RGB")
        return pytesseract.image_to_string(img)
    except Exception as e:
        print(f"Image extract error: {e}")
        return ""

# Interactive document chat function
def chat_about_document(text, question):
    """
    Interactive chat about the document
    
    Args:
        text (str): Document text
        question (str): User's question about the document
    
    Returns:
        str: AI-generated answer
    """
    # If Gemini is not available, return a fallback response
    if not GEMINI_AVAILABLE:
        return "Chat functionality requires Gemini API integration."
    
    prompt = f"""
    Based on the following document, answer the question accurately and concisely.
    
    Document:
    {text[:10000]}  # Limit for context window
    
    Question: {question}
    
    Answer:
    """
    
    try:
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Unable to answer the question due to: {str(e)}"

# Routes
@app.route("/active", methods=["GET"])
def active():
    return "active"

@app.route("/uploads", methods=["POST"])
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    filename = file.filename.lower()
    text = ""

    if filename.endswith(".pdf"):
        text = extract_pdf(file.stream)
    elif filename.endswith(".docx"):
        text = extract_docx(file.stream)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        text = extract_image(file.stream)
    else:
        return jsonify({"error": "Unsupported file type"}), 400

    is_ok, details = classify_agreement(text)

    if not is_ok:
        return jsonify({
            "error": "Rejected: Not a valid agreement.",
            "details": details
        }), 400

    # Perform enhanced analysis
    analysis = analyze_legal_document(text)
    
    return jsonify({
        "filename": file.filename,
        "extracted_text": text,
        "analysis": analysis,
        "timestamp": datetime.now().isoformat()
    })

@app.route("/enhanced_analysis", methods=["POST"])
def enhanced_document_analysis():
    """
    Enhanced document analysis endpoint
    """
    try:
        print("Received request to /enhanced_analysis")
        if "file" not in request.files:
            print("No file uploaded in request")
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files["file"]
        print(f"Received file: {file.filename}")
        if file.filename == "":
            print("No file selected")
            return jsonify({"error": "No file selected"}), 400

        # Extract text (using existing functions)
        filename = file.filename.lower()
        print(f"Processing file type: {filename}")
        text = ""
        
        if filename.endswith(".pdf"):
            print("Extracting PDF...")
            text = extract_pdf(file.stream)
        elif filename.endswith(".docx"):
            print("Extracting DOCX...")
            text = extract_docx(file.stream)
        elif filename.endswith((".png", ".jpg", ".jpeg")):
            print("Extracting image...")
            text = extract_image(file.stream)
        else:
            print(f"Unsupported file type: {filename}")
            return jsonify({"error": "Unsupported file type"}), 400
        
        print(f"Extracted text length: {len(text)}")
        if len(text) == 0:
            print("Warning: No text extracted from file")
            return jsonify({
                "error": "Could not extract text from file",
                "filename": file.filename
            }), 400
        
        # Check if it's a valid agreement (using existing function)
        print("Classifying agreement...")
        is_ok, details = classify_agreement(text)
        print(f"Classification result: {is_ok}, details: {details}")
        if not is_ok:
            return jsonify({
                "error": "Rejected: Not a valid agreement.",
                "details": details
            }), 400
        
        # Perform enhanced analysis
        print("Performing enhanced analysis...")
        try:
            analysis = analyze_legal_document(text)
            print(f"Analysis completed successfully")
        except Exception as analysis_error:
            print(f"Error in analyze_legal_document: {analysis_error}")
            import traceback
            traceback.print_exc()
            # Return fallback analysis
            analysis = create_fallback_analysis(text, "unknown")
        
        return jsonify({
            "filename": file.filename,
            "extracted_text": text[:500] + "..." if len(text) > 500 else text,  # Limit text in response
            "analysis": analysis,
            "timestamp": datetime.now().isoformat()
        })
    except Exception as e:
        print(f"Unexpected error in enhanced_document_analysis: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": f"Internal server error: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }), 500

@app.route("/chat", methods=["POST"])
def document_chat():
    """
    Interactive chat about a document
    """
    data = request.get_json()
    document_text = data.get("document_text", "")
    question = data.get("question", "")
    
    if not document_text or not question:
        return jsonify({"error": "Document text and question are required"}), 400
    
    answer = chat_about_document(document_text, question)
    
    return jsonify({
        "question": question,
        "answer": answer,
        "timestamp": datetime.now().isoformat()
    })

@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    text = request.form.get("text", "")
    output = io.BytesIO()
    doc = SimpleDocTemplate(output)
    styles = getSampleStyleSheet()
    story = [Paragraph(line, styles["Normal"]) for line in text.split("\n")]
    doc.build(story)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.pdf")

@app.route("/export/docx", methods=["POST"])
def export_docx():
    text = request.form.get("text", "")
    output = io.BytesIO()
    d = docx.Document()
    for line in text.split("\n"):
        d.add_paragraph(line)
    d.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.docx")

if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port)
